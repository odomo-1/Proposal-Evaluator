import streamlit as st
import os
import tempfile
import docx
from docx import Document
import re
from io import BytesIO
from spellchecker import SpellChecker

# --- Constants ---
STANDARD_SECTIONS = [
    "Table of content",
    "Introduction",
    "Background",
    "Objective",
    "Methodology" or "Approach",
    "Project Team",
    "About Sahel",
    "Budget",
    "Work Plan",
]

# --- Helpers ---
def extract_text(file):
    text = ""
    if file.name.endswith('.docx'):
        temp_path = os.path.join(tempfile.gettempdir(), file.name)
        with open(temp_path, 'wb') as f:
            f.write(file.read())
        doc = Document(temp_path)
        for para in doc.paragraphs:
            text += para.text + '\n'
    return text

def evaluate_proposal(text, required_sections, doc):
    lower_text = text.lower()

    # Section Presence Check (search for key phrases)
    section_results = {}
    for sec in required_sections:
        found = any(sec.lower() in para.text.lower() for para in doc.paragraphs)
        section_results[sec] = found

    section_score = sum(section_results.values())
    section_percentage = (section_score / len(required_sections)) * 100

    # Formatting & Presentation Check
    formatting_results = formatting_check(doc)

    # Total score calculation considering all criteria
    total_score = 0
    max_score = 4  # 4 main evaluation criteria: sections, font, font size, spelling issues

    # Section presence is 50% of the overall score
    total_score += section_percentage * 0.50

    # Spelling issues check is 25% of the overall score
    spelling_score = 0
    if len(formatting_results['spelling_issues']) == 0:
        spelling_score = 100
    else:
        spelling_score = max(0, 100 - len(formatting_results['spelling_issues']) * 10)
    total_score += spelling_score * 0.25

    # Font style and size check is 25% of the overall score
    font_style_score = 100 if formatting_results['font_ok'] else 0
    font_size_score = 100 if formatting_results['font_size_ok'] else 0
    total_score += (font_style_score + font_size_score) * 0.25

    # Round the total score to the nearest whole number
    total_score = round(total_score)

    # Recommendations
    missing_sections = [sec for sec, present in section_results.items() if not present]
    recommendations = []
    if missing_sections:
        recommendations.append(f"Missing sections: {', '.join(missing_sections)}")

    if formatting_results['spelling_issues']:
        recommendations.append("Spelling issues found in the document.")

    if not formatting_results['font_ok']:
        recommendations.append("Document should use font 'Tenorite' throughout.")

    if not formatting_results['font_size_ok']:
        recommendations.append("Body text should use font size 11.")

    return {
        'sections': section_results,
        'score': total_score,
        'recommendations': recommendations,
        'formatting': formatting_results
    }

def formatting_check(doc):
    spell = SpellChecker()
    text = "\n".join([para.text for para in doc.paragraphs])
    words = re.findall(r'\b\w+\b', text.lower())
    misspelled = spell.unknown(words)
    spelling_issues = list(misspelled)[:15]  # Show up to 15 misspelled words

    # Check for font style "Tenorite" and font size 11 in body text
    font_ok = True
    font_size_ok = True
    for para in doc.paragraphs:
        for run in para.runs:
            # Ensure that the font is Tenorite in both body and heading text
            if run.font.name and run.font.name.lower() != "tenorite":
                font_ok = False
            # Check if font size is 11 for body text (not for headings)
            if run.font.size and run.font.size.pt != 11:
                if para.style.name != 'Heading 1' and para.style.name != 'Heading 2' and para.style.name != 'Heading 3':
                    font_size_ok = False

        if not font_ok or not font_size_ok:
            break

    return {
        "spelling_issues": spelling_issues,
        "font_ok": font_ok,
        "font_size_ok": font_size_ok
    }

def create_word_report(evaluation):
    doc = Document()
    doc.add_heading("Proposal Evaluation Report", level=1)

    doc.add_heading("Section Check", level=2)
    for section, found in evaluation['sections'].items():
        doc.add_paragraph(f"{section}: {'Present' if found else 'Missing'}")

    doc.add_heading("Formatting & Presentation", level=2)

    if evaluation['formatting']['spelling_issues']:
        doc.add_paragraph("Spelling Issues Detected:")
        doc.add_paragraph(", ".join(evaluation['formatting']['spelling_issues']))
    else:
        doc.add_paragraph("No major spelling issues detected.")

    if evaluation['formatting']['font_ok'] and evaluation['formatting']['font_size_ok']:
        doc.add_paragraph("Font style and size meet organizational standards (Tenorite, size 11).")
    else:
        doc.add_paragraph("Font style does not match standard (Tenorite) or font size is not 11 in body text.")

    doc.add_heading("Overall Score", level=2)
    doc.add_paragraph(f"{evaluation['score']}%")

    doc.add_heading("Recommendations", level=2)
    if evaluation['recommendations']:
        for rec in evaluation['recommendations']:
            doc.add_paragraph(f"- {rec}")
    else:
        doc.add_paragraph("All criteria met. Great job!")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit UI ---
st.title("Proposal Evaluator (DOCX Only)")

uploaded_proposal = st.file_uploader("Upload Proposal (.docx only)", type=["docx"])

evaluation = None
if uploaded_proposal and st.button("Evaluate Proposal"):
    st.success("Proposal uploaded successfully.")

    prop_text = extract_text(uploaded_proposal)
    doc = Document(uploaded_proposal)

    with st.spinner("Evaluating proposal..."):
        evaluation = evaluate_proposal(prop_text, STANDARD_SECTIONS, doc)

if evaluation:
    st.subheader("Evaluation Results")

    st.write("### Section Check")
    for section, found in evaluation['sections'].items():
        st.write(f"- **{section}**: {'✅' if found else '❌'}")

    st.write("### Formatting & Presentation")

    if evaluation['formatting']['spelling_issues']:
        st.warning("Spelling Issues Detected:")
        st.write(", ".join(evaluation['formatting']['spelling_issues']))
    else:
        st.success("No major spelling issues detected.")

    if evaluation['formatting']['font_ok'] and evaluation['formatting']['font_size_ok']:
        st.success("Font style and size meet organizational standards (Tenorite, size 11).")
    else:
        st.warning("Font style does not match standard (Tenorite) or font size is not 11 in body text.")

    st.write(f"### Overall Score: **{evaluation['score']}%**")

    st.write("### Recommendations")
    if evaluation['recommendations']:
        for rec in evaluation['recommendations']:
            st.warning(rec)
    else:
        st.success("Your proposal aligns well with the standards!")

    word_buffer = create_word_report(evaluation)
    st.download_button(
        label="Download Evaluation Report (.docx)",
        data=word_buffer,
        file_name="proposal_evaluation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
